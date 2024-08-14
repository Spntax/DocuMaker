from tkinter import *
from tkinter import ttk 
from tkinter import filedialog
from tkinter import messagebox
from docx import Document
from pdfconvert import convert
import msoffice2pdf
import comtypes.client
import docx
import shutil
import DocData
import os
import subprocess
import DB
import Table
import time
import webbrowser
print("hehe haha")

# Global variables ˇˇˇˇˇˇˇˇˇˇˇˇˇˇˇˇˇˇˇˇˇˇˇˇˇˇˇˇˇˇˇˇˇˇˇˇˇˇˇˇ
main_window = Tk()

template_path = "bin/Base_document.docx"
output_path = 'bin/Output_fun.docx'
outputPdf_path = "bin/Output.pdf"

doc = Document(template_path)

documentData = DocData.DocumentData()
usersDB = DB.UsersDB()
list_data = usersDB.GetWorkTable("")
windowUpdate = False
windowIsAlive = True
currentUserID = 0
tk_name,tk_address,tk_phone,tk_callsign,tk_type,tk_modell = 0,0,0,0,0,0

# Global variables ^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^
def WindowKill():
    global windowIsAlive
    windowIsAlive = False
    print("Crush kill destroy SWAG")

def SaveWorkItem(tk_name,tk_address,tk_phone,tk_note,tk_callsign,tk_type,tk_modell,tk_description,tk_addon,tk_diagnosis):
    global windowUpdate
#    Add Record to Database
    usersDB.InsertRecord(tk_name.get(),tk_address.get(),tk_phone.get(),tk_note.get(),tk_callsign.get(),tk_type.get(),tk_modell.get(),tk_description.get(),tk_addon.get(),tk_diagnosis.get())
#    bbb = Entry(main_window,width=20)
#    bbb.grid(row=main_window.grid_size()[1])
#    bbb.insert(END, "fkjghdfkl")

    windowUpdate = True
    print(windowUpdate)
    
def SearchCustomers():
    def LoadButtonClicked(selectedRecord_):
        global tk_name,tk_address,tk_phone,tk_callsign,tk_type,tk_modell

        beginning=selectedRecord_.index("#")
        userID=selectedRecord_[beginning+1:]
        record = usersDB.GetByUserID(userID)
        print(userID)

        tk_name.delete(0,END)
        tk_address.delete(0,END)
        tk_phone.delete(0,END)
        tk_callsign.delete(0,END)
        tk_type.delete(0,END)
        tk_modell.delete(0,END)

        tk_name.insert(END,str(record[1]))
        tk_address.insert(END,str(record[2]))
        tk_phone.insert(END,str(record[3]))
        tk_callsign.insert(END,str(record[4]))
        tk_type.insert(END,str(record[8]))
        tk_modell.insert(END,str(record[9]))
    def SearchButtonClicked(searchInput_,listbox_):
        listbox_.delete(0,END)
        if searchInput_.isnumeric():
            # Doc ID was searched
            records = usersDB.SearchByDocID(searchInput_)
            for record in records:
                fullText = "Sorszám:"+str(record[13])+", "+str(record[1])+", "+str(record[2])+", "+str(record[3])+", "+str(record[4])+", "+str(record[8])+", "+str(record[9])+",                             #"+str(record[0])
                listbox_.insert(END,fullText)       
        else:
            # Name was searched
            records = usersDB.SearchByName(searchInput_)
            for record in records:
                fullText = "Sorszám:"+str(record[13])+", "+str(record[1])+", "+str(record[2])+", "+str(record[3])+", "+str(record[4])+", "+str(record[8])+", "+str(record[9])+",                             #"+str(record[0])
                listbox_.insert(END,fullText)
    customerSearchWindow = Tk()
    customerSearchWindow.eval('tk::PlaceWindow . center')
    customerSearchWindow.title("Keresés")
    bottomframe = Frame(customerSearchWindow)
    bottomframe.pack( side = BOTTOM )
    bottomerframe = Frame(bottomframe)
    bottomerframe.pack( side = BOTTOM )
    Label(customerSearchWindow, text='Név/Sorszám', anchor='w').pack(fill='both',padx=10,pady=(10,0))
    tk_entry = Entry(customerSearchWindow,width=75)
    tk_entry.pack(padx=(10,10), pady=(0,10), side=LEFT)
    tk_searchButton = Button(customerSearchWindow, text="Keresés", command=lambda:(SearchButtonClicked(tk_entry.get(),tk_listbox)))
    tk_searchButton.pack(padx=(0,10),pady=(0,10), side=LEFT)
    tk_listbox = Listbox(bottomframe,width=150,height=15)
    tk_listbox.pack(padx=10,pady=(20,0))
    tk_loadButton = Button(bottomerframe,text="Betöltés", command=lambda:(LoadButtonClicked(tk_listbox.get(ACTIVE))), width=15).pack(padx=(800,10),pady=(10,20))
    
# Menu Functions here -------------------------------------
def UpdateState(ID_,newState_):
    global windowUpdate
    usersDB.UpdateState(ID_,newState_)
    windowUpdate=True

def DeleteWorkItem(ID_,newState_):
    def YesClicked(ID_,newState_):
        global windowUpdate
        usersDB.UpdateState(ID_,newState_)
        delete_popup.destroy()
        windowUpdate=True
    delete_popup=Tk()
    delete_popup.title("Figyelmeztetés")
    delete_popup.eval('tk::PlaceWindow . center')
    Label(delete_popup,text="Biztos hogy törölni szeretnéd ezt a bejegyzést?",font=('Arial',11,'bold')).grid(row=0,column=0,columnspan=2, padx=10,pady=20)
    Button(delete_popup, text="Igen", width=10, command=lambda: YesClicked(ID_,newState_)).grid(row=1,column=0,padx=(70,20),pady=20)
    Button(delete_popup, text='Nem', width=10, command = lambda: delete_popup.destroy()).grid(row=1,column=1,padx=(20,70),pady=20)


def UpdateComment(ID_):
    def CommitCommentUpdate(newComment):
        global windowUpdate
        usersDB.UpdateComment(ID_,newComment)
        windowUpdate=True
        commentUpdateWindow.destroy()
    currentComment = usersDB.GetComment(ID_)
    commentUpdateWindow = Tk()
    commentUpdateWindow.eval('tk::PlaceWindow . center')
    commentUpdateWindow.title("Komment")
    #tk_curID = Label(commentUpdateWindow, text="Komment", font=('Arial',11,'bold'))
    #tk_curID.pack(padx=10, pady=(10,0))
    #Label(commentUpdateWindow, text="Adj meg új dokument sorszámot:").pack(padx=10, pady=(25,0))
    tk_Comment = Entry(commentUpdateWindow,width=40)
    tk_Comment.pack(padx=10)
    tk_Comment.insert(0,currentComment)
    Button(commentUpdateWindow, text="OK", width=10, command=lambda: CommitCommentUpdate(tk_Comment.get())).pack(padx=100,pady=10)

def PrintDocument(userID):
    global windowUpdate

    doc = Document(template_path)    
    recordToExport = usersDB.GetByUserID(userID)
    documentID = usersDB.GetDocumentID(True)
    usersDB.UpdateRecordDocID(userID,documentID)
    documentData.ExtractDataFromDB(recordToExport,documentID)

# Update the template Docx file and replace the data
    documentData.UpdateData(doc)
    doc.save(output_path)
    #shutil.copyfile(output_path, "bin/Output.docx")
    sourcePath = os.path.abspath("bin/Output_fun.docx")
    outputPath = os.path.abspath("bin/")
    result = convert(source=sourcePath, output_dir=outputPath, soft=1)
    print(result)

    #pdf_format = 17
    #word = comtypes.client.CreateObject("Word.Application")
    #in_file = word.Documents.Open(output_path)
    #in_file.SaveAs(os.path.abspath(outputPdf_path),FileFormat=pdf_format)
    #in_file.Close()
    os.system('start bin/alt_Output_fun.Pdf')

def OpenTemplateDoc():
    os.system('start bin/Base_document.docx')

def CheckUpdates():
    webbrowser.open("https://github.com/Spntax/DocuMaker/releases")

def SaveDB():
    filename = filedialog.asksaveasfilename(initialdir = "/", title = "Select a File", filetypes = (("Sqlite DB","*.db*"), ("all files","*.*")))
    shutil.copyfile("bin/users_db.db",filename+".db")

def ImportDB():
    global windowUpdate
    filename = filedialog.askopenfilename(initialdir = "/", title = "Select a File", filetypes = (("Sqlite DB","*.db*"), ("all files","*.*")))
    shutil.copyfile("bin/users_db.db","bin/users_db_old.db")
    shutil.copyfile(filename,"bin/users_db.db")
    windowUpdate = True

def DeleteDB():
    msg_box = messagebox.askquestion(
        "Adatbázis Törlése",
        "Biztos hogy törölni akarod az adatbázist? \n\nEgy biztonsági mentés fog készülni a jelenlegi adatbázisról.",
        icon="warning",
    )
    if msg_box =="yes":
        global windowUpdate
        shutil.copyfile("bin/users_db.db","bin/users_db_deleted.db")
        shutil.copyfile("bin/users_db_empty.db","bin/users_db.db")
        windowUpdate = True
    

def SetDocID():
    def updateDocID(newID):
        usersDB.UpdateDocID(newID)
        tk_curID.config(text=newID)
        tk_DocID.delete(0,END)
    docIDUpdateWindow = Tk()
    docIDUpdateWindow.eval('tk::PlaceWindow . center')
    docIDUpdateWindow.title("Dokument Sorszám")
    tk_curID = Label(docIDUpdateWindow, text=""+str(usersDB.GetDocumentID(False))+"", font=('Arial',16,'bold'))
    tk_curID.pack(padx=10, pady=(10,0))
    Label(docIDUpdateWindow, text="Adj meg új dokument sorszámot:").pack(padx=10, pady=(25,0))
    tk_DocID = Entry(docIDUpdateWindow)
    tk_DocID.pack(padx=10)
    Button(docIDUpdateWindow, text="OK", width=10, command=lambda: updateDocID(tk_DocID.get())).pack(padx=100,pady=10)

# Menu Functions here -------------------------------------

def CreateWorkItemTable(root,list_data):
    def do_popup(event,userID): 
        try: 
            m.tk_popup(event.x_root, event.y_root)
            global currentUserID
            currentUserID = userID
        finally: 
            m.grab_release() 
    if len(list_data)>0:
        total_rows = len(list_data)
        total_columns = len(list_data[0])
        total_columns = total_columns-1
        # code for creating table
        for j in range(total_columns):
            if j!= 0:
                for i in range(total_rows):
                    # Check if item is not "removed"
                    if(list_data[i][total_columns]!=3):
                        m = Menu(root, tearoff=0)
                        #m.add_command(label="Komment")
                        m.add_command(label="Állapot 1", command=lambda i_=i:UpdateState(currentUserID,0))
                        m.add_command(label="Állapot 2", command=lambda i_=i:UpdateState(currentUserID,1))
                        m.add_command(label="Állapot 3", command=lambda i_=i:UpdateState(currentUserID,2))
                        m.add_separator()
                        m.add_command(label="Nyomtatás", command=lambda i_=i:PrintDocument(currentUserID))
                        m.add_separator()
                        m.add_command(label="Komment", command=lambda i_=i:UpdateComment(currentUserID))
                        m.add_separator()
                        m.add_command(label="Törlés", command=lambda i_=i:DeleteWorkItem(currentUserID,3))
                        e = Entry(root, width=20, fg='blue',
                                font=('Arial',11))
                        if(list_data[i][total_columns]==1):
                            e.configure({"disabledbackground": "yellow"})
                        if(list_data[i][total_columns]==2):
                            e.configure({"disabledbackground": "lawngreen"})

                        e.grid(row=i+23, column=j-1)
                        e.insert(END, list_data[i][j])
                        e.configure(state=DISABLED)
                        e.bind("<Button-3>",lambda event, user_id=list_data[i][0]: do_popup(event, user_id))

def DrawMainWindow():
    global tk_name,tk_address,tk_phone,tk_callsign,tk_type,tk_modell
    def NameSearch(e):
        position = tk_name.index(INSERT)
        #print(position)
        #print(e.char)
        if e.char == "":
            print("yahooo")
        else:
            predictedText = usersDB.GetName(tk_name.get())
            #print(predictedText)
            tk_name.insert(position,predictedText[position:])
            tk_name.selection_range(position,END)
    def NameSelectedFromList(ID_,nameSelectWindow):
        end = ID_.index(",")
        choosenUserData = usersDB.GetByUserID(ID_[1:end])
        tk_name.delete(0,END)
        tk_address.delete(0,END)
        tk_phone.delete(0,END)
        tk_name.insert(0,choosenUserData[1])
        tk_address.insert(0,choosenUserData[2])
        tk_phone.insert(0,choosenUserData[3])
        nameSelectWindow.destroy()
    def NameSelected(e):
        name = tk_name.get()
        records = usersDB.SearchByName(name.title())
        notedRecords = []
        filteredRecords = []
        counter = 0
        print("Lenght of the records:"+str(len(records)))
        for record in records:
            isDuplicate = False
            print("Now checking record counter:"+str(counter))
            for name_,address_ in notedRecords:
                print("Noted name:"+str(name_)+", Noted address:"+str(address_)+",  Record name:"+str(record[1]+", Record address:"+str(record[2])))
                if record[1] == name_ and record[2] == address_:
                    print("Match was found")
                    print("popping this record:"+str(records[counter]))
                    isDuplicate = True
                    break
            if not isDuplicate:
                print("Match was not found")
                notedRecords.append((record[1],record[2]))
                filteredRecords.append(records[counter])
            counter = counter+1
            
        records = filteredRecords
        print("Len of records:")
        print(len(records))
        if len(records) <= 1:
            tk_name.delete(0,END)
            tk_address.delete(0,END)
            tk_phone.delete(0,END)
            tk_name.insert(0,records[0][1])
            tk_address.insert(0,records[0][2])
            tk_phone.insert(0,records[0][3])
        else:
            nameSelectWindow = Tk()
            Label(nameSelectWindow, text="Több azonos nevű személy van elmentve az adatbázisban \n Válassz a listából").pack(padx=10, pady=(10,0))
            tk_listbox = Listbox(nameSelectWindow,width=100,height=5)
            tk_listbox.pack(padx=10,pady=20)
            tk_okButton = Button(nameSelectWindow, text = "Ok", width=20, command= lambda: NameSelectedFromList(tk_listbox.get(ACTIVE),nameSelectWindow))
            tk_okButton.pack(padx=10, pady=10)
            for record in records:
                fullText = "#"+str(record[0])+", "+str(record[1])+", "+str(record[2])+", "+str(record[3])
                tk_listbox.insert(END,fullText)

    MenuBar = Menu(main_window)
    # File tab ----------
    FileMenu = Menu(MenuBar, tearoff=0)
    MenuBar.add_cascade(label="File", menu = FileMenu)
    FileMenu.add_command(label="Minta Dokumentum Szerkesztése", command=OpenTemplateDoc)
    FileMenu.add_command(label="Dokumentum sorszám", command=SetDocID)
    # Database tab--------------
    DBMenu = Menu(MenuBar, tearoff=0)
    MenuBar.add_cascade(label="Adatbázis", menu = DBMenu)
    DBMenu.add_command(label="Adatbázis Mentése", command=SaveDB)
    DBMenu.add_command(label="Adatbázis Betöltése", command=ImportDB)
    DBMenu.add_separator()
    DBMenu.add_command(label="Adatbázis Törlése", command=DeleteDB)
    # Update tab--------------
    UpdateMenu = Menu(MenuBar, tearoff=0)
    MenuBar.add_cascade(label="Frissítés", menu = UpdateMenu)
    UpdateMenu.add_command(label="Verzíok megnyitása", command=CheckUpdates)
    Label(main_window, text='Megrendelő Neve:').grid(row=0, column=0,sticky="e")
    tk_name = Entry(main_window)
    tk_name.grid(row=0, column=1, columnspan=2, sticky="EW", padx=(20,80))
    tk_name.bind("<KeyRelease>",NameSearch)
    tk_name.bind("<Return>",NameSelected)
    Label(main_window, text='Cím').grid(row=1,sticky="e")
    tk_address = Entry(main_window)
    tk_address.grid(row=1,column=1,columnspan=2, sticky="EW", padx=(20,20))
    Label(main_window, text='Telefon').grid(row=2,sticky="e")
    tk_phone = Entry(main_window)
    tk_phone.grid(row=2,column=1,columnspan=2, sticky="EW", padx=(20,130))
    Label(main_window, text='Megjegyzés').grid(row=3,sticky="e")
    tk_note = Entry(main_window)
    tk_note.grid(row=3,column=1)
    Label(main_window, text='Megnevezés').grid(row=0,column=2,columnspan=2,sticky="e",padx=(40,45))
    tk_callsign = Entry(main_window)
    tk_callsign.grid(row=0,column=3,columnspan=2, sticky="EW", padx=(130,20))
    Label(main_window, text='Típus').grid(row=1,column=2,columnspan=2,sticky="e",padx=(40,45))
    tk_type = Entry(main_window)
    tk_type.grid(row=1,column=3,columnspan=2, sticky="EW", padx=(130,20))
    Label(main_window, text='Modell').grid(row=2,column=2,columnspan=2,sticky="e",padx=(40,45))
    tk_modell = Entry(main_window)
    tk_modell.grid(row=2,column=3,columnspan=2, sticky="EW", padx=(130,20))
    # Blank line filler
    Label(main_window, text='').grid(row=4,column=0)
    Label(main_window, text='Hibajelenség').grid(row=5,column=0,sticky="e")
    tk_description = Entry(main_window)
    tk_description.grid(row=5,column=1, columnspan=2 ,sticky="EW",padx=(20,100))
    Label(main_window, text='Tartozékok').grid(row=6,column=0,sticky="e")
    tk_addon = Entry(main_window)
    tk_addon.grid(row=6,column=1,columnspan=2,sticky="EW",padx=(20,100))
    Label(main_window, text='Diagnózis',).grid(row=5,column=2,columnspan=2,sticky="e",padx=(40,110))
    tk_diagnosis = Entry(main_window)
    tk_diagnosis.grid(row=5,column=3,columnspan=2, sticky="EW", padx=(70,20))

    Button(main_window, text='Keresés', width=15, command=lambda: SearchCustomers()).grid(row=19,column=0,columnspan=3,pady=(25,1))
    Button(main_window, text='Mentés', width=15, command=lambda: SaveWorkItem(tk_name,tk_address,tk_phone,tk_note,tk_callsign,tk_type,tk_modell,tk_description,tk_addon,tk_diagnosis)).grid(row=19,column=3,columnspan=1,pady=(25,1))

    # ----------------------------- Workitem list --------------------------------
    sep = ttk.Separator(main_window, orient="horizontal").grid(row=20, columnspan=999,pady=15,padx=20,sticky="ew")
    Label(main_window,text="Felvett rendelések", font=('Arial',16,'bold')).grid(row=21,column=0,columnspan=5)
    Label(main_window,text="Név").grid(row=22,column=0)
    Label(main_window,text="Lakhely").grid(row=22,column=1)
    Label(main_window,text="Telefonszám").grid(row=22,column=2)
    Label(main_window,text="Eszköz").grid(row=22,column=3)
    Label(main_window,text="Komment").grid(row=22,column=4)

    main_window.config(menu=MenuBar)

    CreateWorkItemTable(main_window,list_data)

    main_window.title("DocuMaker 1.1")
    main_window.eval('tk::PlaceWindow . center')
    icon = PhotoImage(file="bin/icon.png")
    main_window.iconphoto(icon,icon)

DrawMainWindow()
main_window.protocol("WM_DELETE_WINDOW",  WindowKill)

windowIsAlive = True

# Main loop start here
while windowIsAlive:
    main_window.update()
    main_window.update_idletasks()
    #print(windowUpdate)
    if windowUpdate:
        windowUpdate = False
        main_window.destroy()
        main_window = Tk()
        main_window.protocol("WM_DELETE_WINDOW",  WindowKill)
        windowIsAlive = True
        list_data = usersDB.GetWorkTable("")
        print("yippie")
        DrawMainWindow()
        main_window.update()